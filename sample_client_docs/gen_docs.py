"""Generate sample client documents for StackForge pipeline testing."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import csv, json, os

OUT = "c:/StackForge/sample_client_docs"

CLIENT = "MediConnect Health Systems"
VENDOR = "Stackular Technologies"
PROJECT = "CareFlow — Patient & Telehealth Management Platform"

# ─── DOCX helpers ────────────────────────────────────────────────────────────

def heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def table_2col(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        t.cell(i, 0).text = k
        t.cell(i, 1).text = v
    doc.add_paragraph()

# ─── PDF helpers ─────────────────────────────────────────────────────────────

def pdf_doc(filename, title):
    path = os.path.join(OUT, filename)
    doc = SimpleDocTemplate(path, pagesize=letter,
                            rightMargin=inch, leftMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("SectionHead", parent=styles["Heading2"],
                               textColor=colors.HexColor("#1F4E79"), spaceAfter=6))
    styles.add(ParagraphStyle("Body2", parent=styles["Normal"],
                               fontSize=10, leading=14, alignment=TA_JUSTIFY))
    styles.add(ParagraphStyle("TitlePage", parent=styles["Title"],
                               fontSize=22, textColor=colors.HexColor("#1F4E79")))
    return doc, styles

def pdf_h(styles, text, level="SectionHead"):
    return Paragraph(text, styles[level])

def pdf_p(styles, text):
    return Paragraph(text, styles["Body2"])

def pdf_sp(n=12):
    return Spacer(1, n)

def pdf_table(data, col_widths=None):
    t = Table(data, colWidths=col_widths)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1F4E79")),
        ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
        ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0), (-1,-1), 9),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#EBF3FB")]),
        ("GRID",       (0,0), (-1,-1), 0.5, colors.HexColor("#BDD7EE")),
        ("VALIGN",     (0,0), (-1,-1), "TOP"),
        ("PADDING",    (0,0), (-1,-1), 5),
    ]))
    return t

# =============================================================================
# 1. SOW — .docx
# =============================================================================

def gen_sow_docx():
    doc = Document()
    doc.core_properties.author = "MediConnect Procurement"
    doc.core_properties.title = "Statement of Work — CareFlow"

    heading(doc, "STATEMENT OF WORK", 1)
    heading(doc, "CareFlow — Patient & Telehealth Management Platform", 2)
    doc.add_paragraph()

    table_2col(doc, [
        ("Client",         CLIENT),
        ("Vendor",         VENDOR),
        ("Document ID",    "SOW-MC-2026-001"),
        ("Version",        "1.4"),
        ("Effective Date", "June 10, 2026"),
        ("Expiry Date",    "June 9, 2027"),
        ("Contract Value", "USD $1,850,000"),
        ("Prepared By",    "Sarah Chen, VP Engineering — MediConnect"),
        ("Approved By",    "Dr. Rajiv Patel, CTO — MediConnect"),
    ])

    heading(doc, "1. Background & Purpose", 2)
    para(doc, "MediConnect Health Systems operates a network of 47 outpatient clinics across the Pacific Northwest. The organisation currently manages patient scheduling, clinical documentation, and provider communications across three disconnected legacy systems (Athena Health v14, an in-house PHP scheduling portal, and a Microsoft Teams-based telehealth workaround). The resulting fragmentation causes an estimated 3.2 hours of administrative overhead per provider per day, appointment no-show rates of 22%, and recurring HIPAA audit findings related to unsecured messaging.\n\nThe CareFlow platform will consolidate all patient-facing and provider-facing workflows into a single HIPAA-compliant, cloud-native SaaS application. This SOW defines the scope of work Stackular Technologies will deliver under Contract MC-2026-TECH-04.")

    heading(doc, "2. Objectives", 2)
    for obj in [
        "Reduce per-provider administrative overhead by ≥40% within 6 months of go-live.",
        "Reduce appointment no-show rate from 22% to ≤10% via automated SMS/email reminders.",
        "Achieve HIPAA compliance and pass third-party audit within 90 days of go-live.",
        "Eliminate reliance on Microsoft Teams for clinical communications.",
        "Provide real-time patient vitals dashboard for providers during telehealth sessions.",
        "Enable integration with existing EHR (Athena Health API v2) within Phase 2.",
    ]:
        doc.add_paragraph(obj, style="List Bullet")

    heading(doc, "3. Scope of Work", 2)
    heading(doc, "3.1 In Scope", 3)
    for item in [
        "Patient web portal (appointment booking, medical history, secure messaging)",
        "Provider dashboard (schedule management, patient queue, telehealth launcher)",
        "Telehealth video module (WebRTC-based, HIPAA BAA with Twilio)",
        "Automated appointment reminder engine (SMS via Twilio, email via SendGrid)",
        "Role-based access control (Patient, Provider, Admin, Billing, Compliance Officer)",
        "Audit log and compliance reporting module",
        "RESTful API layer for third-party EHR integration",
        "Admin portal (user management, clinic configuration, analytics dashboard)",
        "Mobile applications: iOS and Android (React Native)",
        "Data migration from legacy scheduling system (CSV/API extract)",
        "Staff training (16 hours live + video library)",
        "12-month post-launch support SLA (99.9% uptime guarantee)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    heading(doc, "3.2 Out of Scope", 3)
    for item in [
        "Integration with pharmacy systems (deferred to Phase 3)",
        "AI diagnostic assistance tools",
        "Medical billing / insurance claims processing",
        "Hardware procurement (tablets, kiosk devices)",
        "Legacy Athena Health data migration (client-managed)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    heading(doc, "4. Deliverables", 2)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Deliverable", "Description", "Due Date", "Acceptance Criteria"]):
        t.cell(0, i).text = h
    rows = [
        ("D1 — Discovery Report", "Documented AS-IS process map and gap analysis", "Week 3", "Signed off by CTO"),
        ("D2 — System Architecture", "Cloud architecture diagram, tech stack decision, security model", "Week 6", "Architecture review board approval"),
        ("D3 — Alpha Release", "Core scheduling + provider dashboard functional", "Week 16", "UAT pass rate >= 90%"),
        ("D4 — Beta Release", "All in-scope modules complete, telehealth live", "Week 24", "UAT pass rate >= 95%, zero P1 bugs"),
        ("D5 — Security Audit", "Third-party HIPAA penetration test report", "Week 26", "No critical findings unresolved"),
        ("D6 — Production Launch", "Full go-live with data migration complete", "Week 30", "Uptime >= 99.9% over 7-day burn-in"),
        ("D7 — Training Completion", "All staff trained, documentation delivered", "Week 31", "Training sign-off sheets"),
        ("D8 — Handover Package", "Source code, infra-as-code, runbooks, API docs", "Week 32", "Vendor sign-off"),
    ]
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()

    heading(doc, "5. Timeline & Milestones", 2)
    para(doc, "Total project duration: 32 weeks (June 10, 2026 — January 29, 2027)")
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Light Grid Accent 1"
    for i, h in enumerate(["Phase", "Duration", "Key Activities"]):
        t2.cell(0, i).text = h
    phases = [
        ("Phase 0 — Discovery",       "Weeks 1-4",   "Stakeholder interviews, AS-IS mapping, requirements sign-off"),
        ("Phase 1 — Architecture",    "Weeks 5-7",   "Tech stack finalisation, cloud setup, CI/CD pipeline, security baseline"),
        ("Phase 2 — Core Build",      "Weeks 8-20",  "Patient portal, provider dashboard, scheduling engine, RBAC"),
        ("Phase 3 — Telehealth",      "Weeks 18-24", "Video module, reminder engine, mobile apps (parallel track)"),
        ("Phase 4 — QA & Security",   "Weeks 22-27", "Integration testing, penetration test, UAT, bug remediation"),
        ("Phase 5 — Migration",       "Weeks 26-29", "Data migration dry-run x2, production migration"),
        ("Phase 6 — Launch & Train",  "Weeks 29-32", "Phased go-live (5 pilot clinics -> all 47), staff training"),
    ]
    for row in phases:
        cells = t2.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()

    heading(doc, "6. Payment Schedule", 2)
    t3 = doc.add_table(rows=1, cols=3)
    t3.style = "Light Grid Accent 1"
    for i, h in enumerate(["Milestone", "Amount (USD)", "Due"]):
        t3.cell(0, i).text = h
    payments = [
        ("Contract execution",        "$185,000",  "June 10, 2026"),
        ("D2 — Architecture approved","$277,500",  "July 22, 2026"),
        ("D3 — Alpha UAT passed",     "$370,000",  "October 7, 2026"),
        ("D4 — Beta UAT passed",      "$370,000",  "November 25, 2026"),
        ("D6 — Production launch",    "$462,500",  "January 8, 2027"),
        ("D8 — Handover complete",    "$185,000",  "January 29, 2027"),
    ]
    for row in payments:
        cells = t3.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()
    para(doc, "Total: USD $1,850,000 (exclusive of applicable taxes)")

    heading(doc, "7. Assumptions & Dependencies", 2)
    for a in [
        "MediConnect will provide a dedicated Product Owner (min. 0.5 FTE) for the project duration.",
        "Athena Health API v2 credentials and sandbox access provided by Week 5.",
        "MediConnect IT will provision VPN access for Stackular engineers by Week 2.",
        "HIPAA Business Associate Agreement (BAA) executed with Twilio and AWS before Week 8.",
        "Legacy scheduling data export (CSV) delivered by MediConnect by Week 25.",
        "MediConnect security team available for architecture review in Weeks 5-6.",
    ]:
        doc.add_paragraph(a, style="List Bullet")

    heading(doc, "8. Acceptance Criteria", 2)
    para(doc, "Each deliverable is accepted when: (a) all documented test cases pass at the stated pass rate threshold, (b) no Priority 1 or Priority 2 defects remain open, (c) the MediConnect Product Owner signs the Acceptance Certificate within 5 business days of delivery. Silence beyond 5 business days constitutes deemed acceptance.")

    heading(doc, "9. Change Management", 2)
    para(doc, "Any change to scope, timeline, or cost requires a signed Change Order (CO) before work begins. Changes requested within 72 hours of a milestone do not extend the milestone date unless the CO explicitly states otherwise. Stackular will provide a CO impact assessment within 3 business days of a change request.")

    heading(doc, "10. Warranties & SLAs", 2)
    para(doc, "12-month post-launch warranty covering defect remediation at no charge for P1/P2 bugs. Production SLA: 99.9% monthly uptime (<=43.8 min downtime/month). RTO: 4 hours. RPO: 1 hour. Excluding scheduled maintenance windows (max 4 hours/month, 48 hours notice).")

    heading(doc, "11. Confidentiality & IP", 2)
    para(doc, "All work product, source code, and documentation produced under this SOW is the sole property of MediConnect Health Systems upon full payment. Stackular retains rights to generic, non-client-specific frameworks and tooling. Both parties are bound by the NDA executed February 14, 2026 (NDA-MC-2026-003).")

    heading(doc, "12. Signatures", 2)
    doc.add_paragraph()
    table_2col(doc, [
        ("MediConnect — Authorised Signatory", "Dr. Rajiv Patel, CTO"),
        ("Signature",                          "___________________________"),
        ("Date",                               "___________________________"),
        ("",                                   ""),
        ("Stackular Technologies — Authorised Signatory", "Amir Hosseini, CEO"),
        ("Signature",                          "___________________________"),
        ("Date",                               "___________________________"),
    ])

    doc.save(os.path.join(OUT, "SOW_CareFlow_v1.4.docx"))
    print("  SOW_CareFlow_v1.4.docx")

# =============================================================================
# 2. SOW — PDF
# =============================================================================

def gen_sow_pdf():
    doc, styles = pdf_doc("SOW_CareFlow_v1.4.pdf", "Statement of Work")
    story = []
    story += [pdf_h(styles, "STATEMENT OF WORK", "Title"),
              pdf_h(styles, "CareFlow — Patient & Telehealth Management Platform", "h2"),
              pdf_p(styles, "Client: MediConnect Health Systems  |  Vendor: Stackular Technologies  |  Document ID: SOW-MC-2026-001  |  Version: 1.4  |  Effective: June 10, 2026  |  Value: USD $1,850,000"),
              pdf_sp(20), HRFlowable(width="100%"), pdf_sp(12)]

    story += [pdf_h(styles, "1. Background & Purpose"),
              pdf_p(styles, "MediConnect operates 47 outpatient clinics in the Pacific Northwest. Three disconnected legacy systems generate 3.2 hours of admin overhead per provider per day, a 22% no-show rate, and recurring HIPAA findings. CareFlow consolidates all workflows into a single HIPAA-compliant cloud-native SaaS platform."),
              pdf_sp()]

    meta = [["Field", "Value"],
            ["Client", CLIENT], ["Vendor", VENDOR],
            ["Contract Value", "USD $1,850,000"],
            ["Effective Date", "June 10, 2026"],
            ["Project Duration", "32 weeks"],
            ["Go-Live Target", "January 15, 2027"]]
    story += [pdf_table(meta, [2*inch, 4*inch]), pdf_sp()]

    story += [pdf_h(styles, "2. Deliverables & Payment Schedule")]
    pay = [["Deliverable", "Due", "Amount"],
           ["D1 — Discovery Report", "Week 3", "$185,000"],
           ["D2 — System Architecture", "Week 6", "$277,500"],
           ["D3 — Alpha Release", "Week 16", "$370,000"],
           ["D4 — Beta Release", "Week 24", "$370,000"],
           ["D6 — Production Launch", "Week 30", "$462,500"],
           ["D8 — Handover Package", "Week 32", "$185,000"],
           ["TOTAL", "", "$1,850,000"]]
    story += [pdf_table(pay, [3*inch, 1.2*inch, 1.3*inch]), pdf_sp()]

    story += [pdf_h(styles, "3. SLAs & Warranties"),
              pdf_p(styles, "Production uptime SLA: 99.9% monthly. RTO: 4 hours. RPO: 1 hour. 12-month post-launch warranty on P1/P2 defects at no charge."),
              pdf_sp()]

    story += [pdf_h(styles, "4. Out of Scope"),
              pdf_p(styles, "Pharmacy integration (Phase 3), AI diagnostics, insurance billing, hardware procurement, legacy Athena Health data migration (client-managed)."),
              pdf_sp()]

    doc.build(story)
    print("  SOW_CareFlow_v1.4.pdf")

# =============================================================================
# 3. BRD — .docx
# =============================================================================

def gen_brd_docx():
    doc = Document()
    doc.core_properties.title = "Business Requirements Document — CareFlow"

    heading(doc, "BUSINESS REQUIREMENTS DOCUMENT (BRD)", 1)
    heading(doc, "CareFlow — Patient & Telehealth Management Platform", 2)
    table_2col(doc, [
        ("Document ID",    "BRD-MC-2026-002"),
        ("Version",        "2.1"),
        ("Status",         "APPROVED"),
        ("Date",           "May 28, 2026"),
        ("Author",         "Sarah Chen, VP Engineering"),
        ("Reviewer",       "Dr. Priya Mehta, Chief Medical Officer"),
        ("Approved By",    "Dr. Rajiv Patel, CTO"),
    ])

    heading(doc, "1. Executive Summary", 2)
    para(doc, "MediConnect Health Systems requires a unified digital health platform to replace three aging, disconnected systems. This BRD captures all business requirements gathered during a 6-week discovery phase involving 34 stakeholder interviews, 12 workflow observation sessions, and analysis of 18 months of operational data. Requirements are prioritised using MoSCoW notation (Must / Should / Could / Won't).")

    heading(doc, "2. Business Context", 2)
    heading(doc, "2.1 Problem Statement", 3)
    para(doc, "Administrative staff spend an average of 3.2 hours per provider per day reconciling data across three systems. The 22% no-show rate costs MediConnect approximately $2.3M annually in lost revenue. HIPAA auditors flagged unsecured provider-to-provider messaging via Teams in Q1 2026, resulting in a $45,000 fine and mandatory corrective action plan.")

    heading(doc, "2.2 Business Goals", 3)
    for g in [
        "BG-01: Reduce per-provider administrative overhead by >=40% within 6 months of go-live.",
        "BG-02: Reduce no-show rate from 22% to <=10% via automated reminders.",
        "BG-03: Achieve zero HIPAA findings in the next scheduled audit (Q1 2027).",
        "BG-04: Increase telehealth visit volume by 35% by enabling self-serve booking.",
        "BG-05: Reduce IT operating cost for legacy systems by $180,000/year post-migration.",
    ]:
        doc.add_paragraph(g, style="List Bullet")

    heading(doc, "3. Stakeholder Register", 2)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Stakeholder", "Role", "Interest", "Influence"]):
        t.cell(0, i).text = h
    stakeholders = [
        ("Dr. Rajiv Patel",   "CTO",                    "Platform architecture, cost",       "High"),
        ("Dr. Priya Mehta",   "CMO",                    "Clinical workflows, compliance",    "High"),
        ("Lisa Tanaka",       "Director of Nursing",    "Staff usability, training impact",  "High"),
        ("Marcus Webb",       "Revenue Cycle Manager",  "Billing integration, reporting",    "Medium"),
        ("Carlos Diaz",       "IT Infrastructure Lead", "Cloud hosting, security",           "High"),
        ("Patient Panel (8)", "Patient Representatives","Ease of use, privacy",              "Medium"),
        ("Compliance Officer","HIPAA/Compliance",       "Audit trail, data security",        "High"),
    ]
    for row in stakeholders:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()

    heading(doc, "4. Functional Requirements", 2)

    modules = {
        "4.1 Patient Portal": [
            ("BR-PP-01", "Must",   "Patients shall self-register using email or SSO (Google/Apple)"),
            ("BR-PP-02", "Must",   "Patients shall book, reschedule, and cancel appointments online 24/7"),
            ("BR-PP-03", "Must",   "System shall send appointment reminders at T-48h and T-2h via SMS and email"),
            ("BR-PP-04", "Must",   "Patients shall access their visit history and clinical documents (read-only)"),
            ("BR-PP-05", "Must",   "Patients shall message their care team via HIPAA-secure in-app messaging"),
            ("BR-PP-06", "Should", "Patients shall complete pre-visit intake forms digitally before appointment"),
            ("BR-PP-07", "Should", "Patients shall view and pay outstanding balances"),
            ("BR-PP-08", "Could",  "Patients shall earn loyalty points for on-time attendance"),
        ],
        "4.2 Provider Dashboard": [
            ("BR-PD-01", "Must",   "Providers shall view their daily schedule with patient details in one screen"),
            ("BR-PD-02", "Must",   "Providers shall mark appointment status (arrived, no-show, completed) in real time"),
            ("BR-PD-03", "Must",   "Providers shall initiate a telehealth session directly from the schedule entry"),
            ("BR-PD-04", "Must",   "Providers shall access patient intake forms and visit notes during consultation"),
            ("BR-PD-05", "Should", "Providers shall see patient vitals (if connected wearable) on consultation screen"),
            ("BR-PD-06", "Should", "Providers shall dictate clinical notes (speech-to-text) during consultation"),
            ("BR-PD-07", "Must",   "Providers shall send secure messages to colleagues and care team members"),
        ],
        "4.3 Telehealth Module": [
            ("BR-TH-01", "Must",   "System shall support HD video calls (1080p minimum) between patient and provider"),
            ("BR-TH-02", "Must",   "Video calls shall be encrypted end-to-end (SRTP/DTLS) and HIPAA-compliant"),
            ("BR-TH-03", "Must",   "System shall allow provider to share screen and annotate medical images during call"),
            ("BR-TH-04", "Must",   "System shall record sessions only with explicit patient consent captured digitally"),
            ("BR-TH-05", "Should", "System shall support multi-participant calls (patient + family + specialist)"),
            ("BR-TH-06", "Should", "System shall send SMS/email join link to patient 15 minutes before session"),
            ("BR-TH-07", "Could",  "System shall provide a virtual waiting room with estimated wait time display"),
        ],
        "4.4 Administration & Compliance": [
            ("BR-AC-01", "Must",   "Admins shall manage user accounts, roles, and clinic assignments"),
            ("BR-AC-02", "Must",   "System shall maintain an immutable audit log of all PHI access events"),
            ("BR-AC-03", "Must",   "System shall enforce session timeout after 15 minutes of inactivity"),
            ("BR-AC-04", "Must",   "System shall support MFA for all provider and admin accounts"),
            ("BR-AC-05", "Must",   "System shall generate HIPAA-compliant access reports on demand"),
            ("BR-AC-06", "Should", "Admins shall configure clinic operating hours, provider schedules, and services"),
            ("BR-AC-07", "Must",   "System shall retain PHI data for 7 years in compliance with 45 CFR s164.530"),
        ],
    }

    for mod, reqs in modules.items():
        heading(doc, mod, 3)
        t2 = doc.add_table(rows=1, cols=3)
        t2.style = "Light Grid Accent 1"
        for i, h in enumerate(["ID", "Priority", "Requirement Statement"]):
            t2.cell(0, i).text = h
        for row in reqs:
            cells = t2.add_row().cells
            cells[0].text = row[0]
            cells[1].text = row[1]
            cells[2].text = row[2]
        doc.add_paragraph()

    heading(doc, "5. Non-Functional Requirements", 2)
    nfrs = [
        ("NFR-01", "Performance",  "Patient portal page load time <= 2 seconds at p95 under 10,000 concurrent users"),
        ("NFR-02", "Performance",  "Telehealth call setup latency <= 3 seconds"),
        ("NFR-03", "Availability", "99.9% monthly uptime for all patient-facing services"),
        ("NFR-04", "Security",     "All data at rest encrypted AES-256; all data in transit TLS 1.3+"),
        ("NFR-05", "Security",     "HIPAA Business Associate Agreements in place for all sub-processors"),
        ("NFR-06", "Scalability",  "System shall support 5x peak load without configuration changes (auto-scaling)"),
        ("NFR-07", "Usability",    "Patient portal SUS score >= 80 (usability testing with 12 representative patients)"),
        ("NFR-08", "Compliance",   "SOC 2 Type II certification achieved within 12 months of launch"),
        ("NFR-09", "Accessibility","Patient portal meets WCAG 2.1 AA standard"),
        ("NFR-10", "Data",         "Automated daily backups; point-in-time recovery to within 1 hour"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["ID", "Category", "Requirement"]):
        t.cell(0, i).text = h
    for row in nfrs:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    doc.add_paragraph()

    heading(doc, "6. Constraints", 2)
    for c in [
        "Platform must be hosted on AWS (MediConnect has existing Enterprise Agreement).",
        "Mobile applications must support iOS 15+ and Android 10+.",
        "All patient-facing UI copy must be available in English and Spanish at launch.",
        "Integration with Athena Health API v2 must not require changes to the EHR system.",
        "Budget ceiling: USD $1,850,000 total cost of delivery.",
        "Hard go-live deadline: January 15, 2027 (board-mandated).",
    ]:
        doc.add_paragraph(c, style="List Bullet")

    heading(doc, "7. Acceptance Criteria Summary", 2)
    para(doc, "Acceptance is granted when: (a) all Must requirements pass documented test cases, (b) >=85% of Should requirements pass, (c) SUS score >=80 from patient usability testing, (d) zero open P1 security findings, (e) successful HIPAA pen test report with no critical findings.")

    doc.save(os.path.join(OUT, "BRD_CareFlow_v2.1.docx"))
    print("  BRD_CareFlow_v2.1.docx")

# =============================================================================
# 4. FRS — .docx
# =============================================================================

def gen_frs_docx():
    doc = Document()
    doc.core_properties.title = "Functional Requirements Specification — CareFlow"
    heading(doc, "FUNCTIONAL REQUIREMENTS SPECIFICATION", 1)
    heading(doc, "CareFlow — Patient & Telehealth Management Platform", 2)
    table_2col(doc, [
        ("Document ID", "FRS-MC-2026-003"), ("Version", "1.0"),
        ("Status", "BASELINED"), ("Date", "June 2, 2026"),
        ("Author", "Carlos Diaz / Sarah Chen"), ("System", "CareFlow v1.0"),
    ])

    heading(doc, "1. Introduction", 2)
    para(doc, "This FRS defines the detailed functional behaviour of the CareFlow system. Each requirement is traceable to a BRD business requirement. Test case IDs cross-reference the QA Test Plan (QTP-MC-2026-005).")

    heading(doc, "2. System Overview", 2)
    para(doc, "CareFlow is a multi-tenant SaaS application deployed on AWS. The system comprises: (1) React/TypeScript web application, (2) React Native mobile applications, (3) Node.js/Express API gateway, (4) Python microservices (scheduling engine, reminder engine, analytics), (5) PostgreSQL primary database, (6) Redis cache, (7) Elasticsearch for search and audit, (8) WebRTC telehealth module via Twilio.")

    heading(doc, "3. User Roles & Permissions", 2)
    roles = [
        ("PATIENT",           "Self-register, book/cancel appointments, message care team, join telehealth, view records"),
        ("PROVIDER",          "All Patient access + manage schedule, clinical notes, initiate telehealth, view full patient records"),
        ("CLINIC_ADMIN",      "All Provider access + manage clinic users, configure services, view clinic analytics"),
        ("BILLING",           "Read-only access to appointments and payments; manage billing records"),
        ("COMPLIANCE_OFFICER","Read-only access to all audit logs and compliance reports; cannot modify PHI"),
        ("SUPER_ADMIN",       "Full system access; manage tenants, global configuration, system health"),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.cell(0,0).text = "Role"
    t.cell(0,1).text = "Capabilities"
    for role, caps in roles:
        row = t.add_row().cells
        row[0].text = role
        row[1].text = caps
    doc.add_paragraph()

    heading(doc, "4. Functional Specifications", 2)

    heading(doc, "4.1 Authentication & Session Management", 3)
    specs = [
        ("FS-AUTH-01", "The system shall authenticate patients via email/password or OAuth2 SSO (Google, Apple ID). Password must meet NIST SP 800-63B Level 2: min 8 chars, no complexity rules, breach-check against HaveIBeenPwned API."),
        ("FS-AUTH-02", "All Provider, Admin, and Compliance roles shall require TOTP-based MFA (RFC 6238) at every login. MFA bypass is not permitted under any circumstance."),
        ("FS-AUTH-03", "Sessions shall be invalidated after 15 minutes of inactivity. The system shall display a 2-minute warning before expiry with option to extend."),
        ("FS-AUTH-04", "Refresh tokens shall have a 24-hour TTL for patients, 8-hour TTL for clinical staff. Tokens shall be rotated on each use (sliding window)."),
        ("FS-AUTH-05", "Failed login attempts shall trigger: lockout after 5 failures within 15 minutes; notify account owner by email; require CAPTCHA on next attempt."),
        ("FS-AUTH-06", "All authentication events shall be written to the immutable audit log within 100ms of the event."),
    ]
    for sid, desc in specs:
        para(doc, f"{sid}: {desc}")

    heading(doc, "4.2 Appointment Scheduling Engine", 3)
    specs = [
        ("FS-SCH-01", "The scheduler shall expose provider availability in 15-minute slots. Slot duration is configurable per service type (new patient = 60min, follow-up = 30min, telehealth = 30min)."),
        ("FS-SCH-02", "The system shall prevent double-booking. Slot reservation shall use optimistic locking with a 5-minute hold timer. Expired holds release automatically."),
        ("FS-SCH-03", "The reminder engine shall dispatch SMS (Twilio) and email (SendGrid) at T-48h and T-2h. Failed dispatches shall retry x3 with exponential backoff (60s, 300s, 900s)."),
        ("FS-SCH-04", "Patients shall be able to cancel or reschedule up to T-24h before the appointment without penalty. Within T-24h, cancellation triggers a late-cancel flag on the patient record."),
        ("FS-SCH-05", "The system shall track no-show history and surface a flag on the patient profile if no-show rate exceeds 30% over the preceding 12 months."),
        ("FS-SCH-06", "Scheduling state machine: PENDING -> CONFIRMED -> ARRIVED | NO_SHOW | CANCELLED. Transitions must be timestamped and logged."),
    ]
    for sid, desc in specs:
        para(doc, f"{sid}: {desc}")

    heading(doc, "4.3 Telehealth Video Module", 3)
    specs = [
        ("FS-TH-01", "Video sessions shall be established via Twilio Programmable Video API. Room creation shall be triggered by the provider clicking 'Start Session' <=10 minutes before the scheduled time."),
        ("FS-TH-02", "The system shall display connection quality indicator (green/amber/red) for both participants. If connection drops below 150kbps, the system shall automatically downgrade to audio-only and notify participants."),
        ("FS-TH-03", "Session recording requires explicit patient consent captured as a digital checkbox with timestamp and IP address logged to audit. Recording without consent is a hard block."),
        ("FS-TH-04", "Recordings shall be stored in S3 with server-side AES-256 encryption. Access requires the requesting user's session token and is logged to audit."),
        ("FS-TH-05", "The virtual waiting room shall display: provider name, estimated wait time (+/- 5 min), and a connection test widget. Wait time exceeding 15 minutes triggers an SMS notification to the patient."),
    ]
    for sid, desc in specs:
        para(doc, f"{sid}: {desc}")

    heading(doc, "4.4 Secure Messaging", 3)
    specs = [
        ("FS-MSG-01", "Messages shall be encrypted at rest (AES-256) and in transit (TLS 1.3). Message content shall never be stored in plaintext."),
        ("FS-MSG-02", "Attachments are limited to: JPEG, PNG, PDF, DICOM. Max size 25MB per attachment. DICOM files are rendered via integrated OHIF viewer."),
        ("FS-MSG-03", "Push notifications for new messages shall not include PHI. Notification text: 'You have a new message from [Provider Name].' No message preview."),
        ("FS-MSG-04", "Message threads shall be retained for 7 years per 45 CFR s164.530(j). Deletion by users is not permitted; archival to cold storage after 2 years."),
    ]
    for sid, desc in specs:
        para(doc, f"{sid}: {desc}")

    heading(doc, "4.5 Audit Log", 3)
    specs = [
        ("FS-AUD-01", "All PHI access events (read, create, update, delete) shall generate an immutable audit record containing: user_id, role, action, resource_type, resource_id, timestamp (UTC), IP address, session_id."),
        ("FS-AUD-02", "Audit records shall be written to Elasticsearch with a write-once index policy. No application code shall have DELETE privilege on audit indices."),
        ("FS-AUD-03", "Compliance Officers shall be able to query audit logs by: date range, user, patient, action type, resource type. Query results exportable as CSV."),
        ("FS-AUD-04", "Audit log queries shall return results within 5 seconds for date ranges <=90 days."),
    ]
    for sid, desc in specs:
        para(doc, f"{sid}: {desc}")

    heading(doc, "5. Integration Interfaces", 2)
    integ = [
        ("Athena Health API v2",  "REST",      "Patient demographics sync, visit summary write-back", "Phase 2"),
        ("Twilio Video",          "REST+WebRTC","Telehealth room creation, participant management",    "Phase 1"),
        ("Twilio SMS",            "REST",      "Appointment reminders, MFA codes",                    "Phase 1"),
        ("SendGrid",              "REST",      "Email reminders, account notifications",               "Phase 1"),
        ("AWS S3",                "SDK",       "Recording storage, document storage",                  "Phase 1"),
        ("HaveIBeenPwned API",    "REST",      "Password breach checking at registration",             "Phase 1"),
    ]
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Integration", "Protocol", "Purpose", "Phase"]):
        t.cell(0, i).text = h
    for row in integ:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    doc.add_paragraph()

    doc.save(os.path.join(OUT, "FRS_CareFlow_v1.0.docx"))
    print("  FRS_CareFlow_v1.0.docx")

# =============================================================================
# 5. RFP — PDF
# =============================================================================

def gen_rfp_pdf():
    doc, styles = pdf_doc("RFP_CareFlow_2026.pdf", "Request for Proposal")
    story = []
    story += [pdf_h(styles, "REQUEST FOR PROPOSAL", "Title"),
              pdf_h(styles, "CareFlow — Patient & Telehealth Management Platform", "h2"),
              pdf_p(styles, "Issued by: MediConnect Health Systems  |  RFP Reference: RFP-MC-2026-TECH-01  |  Issue Date: March 15, 2026  |  Response Deadline: April 18, 2026, 17:00 PST"),
              pdf_sp(20), HRFlowable(width="100%"), pdf_sp(12)]

    story += [pdf_h(styles, "1. Overview & Purpose"),
              pdf_p(styles, "MediConnect Health Systems invites qualified technology vendors to submit proposals for the design, development, and deployment of a unified Patient and Telehealth Management Platform (CareFlow). The selected vendor will replace three existing legacy systems and deliver a HIPAA-compliant, cloud-native SaaS solution within a 32-week timeline."),
              pdf_sp()]

    story += [pdf_h(styles, "2. Organisation Background"),
              pdf_p(styles, "MediConnect operates 47 outpatient clinics in Washington, Oregon, and Idaho with approximately 620 licensed providers and 180,000 active patients. Annual revenue: $310M. We process ~850,000 appointments per year. Current systems: Athena Health v14 (EHR), in-house PHP scheduling portal (circa 2014), Microsoft Teams (ad-hoc telehealth). IT team: 23 staff."),
              pdf_sp()]

    story += [pdf_h(styles, "3. Scope Summary")]
    scope_data = [["Module", "Description"],
                  ["Patient Portal", "Self-service booking, records, secure messaging"],
                  ["Provider Dashboard", "Schedule management, patient queue, notes"],
                  ["Telehealth", "WebRTC video, virtual waiting room, recording"],
                  ["Reminder Engine", "Automated SMS/email at T-48h and T-2h"],
                  ["Admin Portal", "User management, analytics, compliance reporting"],
                  ["Mobile Apps", "iOS and Android (React Native)"],
                  ["EHR Integration", "Athena Health API v2 (Phase 2)"],
                  ["Data Migration", "Legacy scheduling data import"]]
    story += [pdf_table(scope_data, [2*inch, 4*inch]), pdf_sp()]

    story += [pdf_h(styles, "4. Mandatory Vendor Requirements"),
              pdf_p(styles, "Proposals failing any of the following will be rejected without further evaluation:")]
    for req in [
        "Minimum 3 years experience delivering HIPAA-compliant healthcare software",
        "At least 2 reference clients with comparable patient portal implementations (>50,000 patients)",
        "AWS-certified architects on the delivery team",
        "Demonstrated SOC 2 Type II certification or active audit in progress",
        "US-based legal entity; all PHI processed on US soil only",
    ]:
        story.append(pdf_p(styles, f"  - {req}"))
    story.append(pdf_sp())

    story += [pdf_h(styles, "5. Evaluation Criteria")]
    eval_data = [["Criterion", "Weight"],
                 ["Technical approach & architecture quality", "25%"],
                 ["Relevant healthcare experience & references", "20%"],
                 ["Security & HIPAA compliance posture", "20%"],
                 ["Proposed timeline & project methodology", "15%"],
                 ["Total cost of ownership", "15%"],
                 ["Team qualifications & staffing plan", "5%"]]
    story += [pdf_table(eval_data, [4*inch, 1.5*inch]), pdf_sp()]

    story += [pdf_h(styles, "6. Submission Instructions"),
              pdf_p(styles, "Submit proposals electronically to procurement@mediconnect.health by April 18, 2026 at 17:00 PST. Subject line: 'CareFlow RFP Response — [Vendor Name]'. Questions must be submitted by April 4, 2026; responses published April 10, 2026."),
              pdf_sp()]

    doc.build(story)
    print("  RFP_CareFlow_2026.pdf")

# =============================================================================
# 6. Project Charter — .docx
# =============================================================================

def gen_charter_docx():
    doc = Document()
    heading(doc, "PROJECT CHARTER", 1)
    heading(doc, "CareFlow — Patient & Telehealth Management Platform", 2)
    table_2col(doc, [
        ("Project Name",      "CareFlow"),
        ("Project Sponsor",   "Dr. Rajiv Patel, CTO — MediConnect"),
        ("Project Manager",   "Sarah Chen, VP Engineering — MediConnect"),
        ("Vendor PM",         "Priya Sharma, Delivery Lead — Stackular Technologies"),
        ("Charter Date",      "June 5, 2026"),
        ("Target Launch",     "January 15, 2027"),
        ("Budget",            "USD $1,850,000"),
        ("Document ID",       "PC-MC-2026-001"),
    ])

    heading(doc, "1. Project Purpose & Justification", 2)
    para(doc, "CareFlow is initiated to eliminate MediConnect's reliance on three disconnected legacy systems, reduce administrative overhead by 40%, cut no-show rates by more than half, and achieve zero HIPAA audit findings by Q1 2027. The project addresses a $2.3M annual revenue leakage from no-shows and a $45,000 Q1 2026 HIPAA fine.")

    heading(doc, "2. Project Objectives", 2)
    for o in [
        "Deliver a production-ready, HIPAA-compliant platform by January 15, 2027.",
        "Achieve >=90% patient portal adoption within 90 days of launch.",
        "Reduce no-show rate to <=10% within 6 months.",
        "Reduce admin overhead per provider per day by >=40% within 6 months.",
        "Pass third-party HIPAA penetration test with zero critical findings.",
    ]:
        doc.add_paragraph(o, style="List Number")

    heading(doc, "3. High-Level Scope", 2)
    para(doc, "In Scope: Patient portal, provider dashboard, telehealth module, reminder engine, admin portal, iOS/Android apps, data migration, staff training, 12-month support.\n\nOut of Scope: Pharmacy integration, AI diagnostics, insurance billing, hardware procurement.")

    heading(doc, "4. Project Organisation", 2)
    org = [
        ("Project Sponsor",         "Dr. Rajiv Patel",   "Strategic decisions, budget approval, risk escalation"),
        ("Project Manager (Client)","Sarah Chen",         "Requirements, UAT, stakeholder alignment"),
        ("Project Manager (Vendor)","Priya Sharma",       "Delivery, team management, status reporting"),
        ("Technical Lead (Vendor)", "Amir Hosseini",      "Architecture, code quality, technical decisions"),
        ("QA Lead (Vendor)",        "Mei Tanaka",         "Test strategy, UAT coordination"),
        ("Security Lead (Vendor)",  "James Okafor",       "Security architecture, pen test management"),
        ("Product Owner (Client)",  "Lisa Tanaka",        "Sprint priorities, feature acceptance"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Role", "Name", "Responsibilities"]):
        t.cell(0, i).text = h
    for row in org:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    doc.add_paragraph()

    heading(doc, "5. Budget Summary", 2)
    table_2col(doc, [
        ("Total Contract Value", "USD $1,850,000"),
        ("Contingency Reserve",  "USD $92,500 (5%) — held by MediConnect"),
        ("Change Order Budget",  "To be approved by Project Sponsor on case-by-case basis"),
        ("Invoicing",            "Milestone-based per SOW-MC-2026-001 Section 6"),
    ])

    heading(doc, "6. Key Risks", 2)
    risks = [
        ("R-01","High",   "Athena Health API v2 credentials delayed","Block Phase 2 integration","Obtain credentials by Week 5; workaround: stub interface"),
        ("R-02","High",   "HIPAA pen test uncovers critical findings","Delay go-live","Weekly security reviews; automated SAST in CI from Week 8"),
        ("R-03","Medium", "Key vendor staff turnover","Knowledge loss, delays","Pair-programming; comprehensive documentation"),
        ("R-04","Medium", "Legacy data quality issues in migration","Corrupted/incomplete records","Dry-run migrations in Weeks 26 and 28; validation scripts"),
        ("R-05","Medium", "Telehealth HIPAA BAA with Twilio delayed","Cannot launch telehealth","Initiate BAA process by Week 4"),
        ("R-06","Low",    "Patient portal adoption below target","No-show goal not met","Onboarding tutorial; clinic champion program"),
    ]
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["ID", "Level", "Risk", "Impact", "Mitigation"]):
        t.cell(0, i).text = h
    for row in risks:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    doc.add_paragraph()

    heading(doc, "7. Success Criteria", 2)
    for s in [
        "No-show rate <=10% within 6 months of launch.",
        "Provider admin overhead reduced by >=40%.",
        "Zero critical or high security findings unresolved at go-live.",
        "Patient portal SUS score >=80.",
        "System availability >=99.9% over first 3 months of operation.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    heading(doc, "8. Charter Approval", 2)
    table_2col(doc, [
        ("Project Sponsor",     "Dr. Rajiv Patel, CTO"),
        ("Signature",           "__________________________"),
        ("Date",                "__________________________"),
        ("",                    ""),
        ("Project Manager",     "Sarah Chen, VP Engineering"),
        ("Signature",           "__________________________"),
        ("Date",                "__________________________"),
    ])

    doc.save(os.path.join(OUT, "ProjectCharter_CareFlow_v1.0.docx"))
    print("  ProjectCharter_CareFlow_v1.0.docx")

# =============================================================================
# 7. Meeting Transcript — .txt
# =============================================================================

def gen_transcript_txt():
    content = """MEETING TRANSCRIPT -- CAREFLOW DISCOVERY WORKSHOP #3
Date: May 14, 2026
Time: 10:00 AM - 12:45 PM PST
Location: MediConnect HQ, Seattle WA -- Boardroom 4B (+ Zoom bridge)
Facilitator: Priya Sharma (Stackular Technologies)
Note-taker: Diego Ruiz (Stackular Technologies)

ATTENDEES:
  MediConnect:
    - Dr. Rajiv Patel (CTO)
    - Sarah Chen (VP Engineering)
    - Lisa Tanaka (Director of Nursing, Eastside Clinic)
    - Marcus Webb (Revenue Cycle Manager)
    - Carlos Diaz (IT Infrastructure Lead)
    - Dr. Emily Rossetti (Provider representative, Portland Clinic)
    - Jordan Kim (Patient Representative, Patient Advisory Board)
  Stackular Technologies:
    - Priya Sharma (Delivery Lead)
    - Amir Hosseini (Technical Lead)
    - Mei Tanaka (QA Lead)
    - Diego Ruiz (Business Analyst)

---------------------------------------------------------------
SECTION 1: SCHEDULING WORKFLOW PAIN POINTS
---------------------------------------------------------------

[10:04] PRIYA SHARMA: Good morning everyone. Today is Workshop 3 -- we are focused
on the scheduling engine and the no-show problem. Lisa, can you walk us through what
a typical day looks like for your front desk staff at Eastside?

[10:05] LISA TANAKA: Sure. So our front desk team of three opens at 7:30 AM. First
thing they do is pull the day's schedule from the PHP portal, cross-reference it
with Athena to check for any insurance changes or demographic updates -- that process
alone takes about 45 minutes every morning for the three of them combined. Then they
start calling patients who did not confirm. We send one automated email when the
appointment is booked, but we have no reminder system after that. No SMS. Nothing
at T-48 or T-2.

[10:08] DR. EMILY ROSSETTI: And from my side as a provider, I get to clinic at 8
and my first patient is usually at 8:30. But I do not know if the first patient is
going to show until they actually walk in. Sometimes I do not find out a patient
no-showed until I walk into the room and it is empty. That is a 30-minute slot gone.

[10:10] MARCUS WEBB: The revenue impact is significant. A 30-min primary care slot
is billed at about $180 on average. We have roughly 35 no-shows per day across the
network. That is $6,300 per day, $1.57M per year in lost revenue just on no-shows.
And we cannot always backfill because we do not have a real-time waitlist.

[10:13] PRIYA SHARMA: Amir, can you speak to what CareFlow's reminder architecture
will look like?

[10:14] AMIR HOSSEINI: Yes. The reminder engine will send two automatic touches --
SMS at 48 hours before the appointment, and again at 2 hours before. Both carry a
one-tap confirm/cancel link. If the patient cancels, the slot is immediately returned
to the available pool and the waitlist engine checks for patients who have requested
that slot or that provider. The whole process is automated. Front desk does not need
to touch it.

[10:16] LISA TANAKA: A real-time waitlist would be huge. Right now we have a paper
waitlist at each clinic. It is embarrassing honestly.

[10:17] CARLOS DIAZ: How does the remind/confirm flow interact with Athena? Because
appointment status changes need to propagate back to the EHR.

[10:19] AMIR HOSSEINI: In Phase 1, CareFlow manages its own appointment state. In
Phase 2, we wire in the Athena v2 webhook so any status change in CareFlow fires an
update to Athena. We will need your Athena API credentials and sandbox access by
Week 5 of the project to scope that integration properly.

[10:21] DR. RAJIV PATEL: That is on us. Carlos, make a note -- we need to get
Stackular the Athena sandbox credentials by June 27th at the latest.

[10:22] CARLOS DIAZ: Noted. I will chase the Athena account manager this week.

---------------------------------------------------------------
SECTION 2: TELEHEALTH REQUIREMENTS
---------------------------------------------------------------

[10:35] PRIYA SHARMA: Let us move to telehealth. Dr. Rossetti, you have been doing
telehealth via Teams for about 18 months now. What are the top three things that
need to be different in CareFlow?

[10:36] DR. EMILY ROSSETTI: Number one -- I need to be able to pull up the patient's
chart during the call without switching windows. Right now I am Alt-Tabbing between
Teams and Athena constantly. It breaks the clinical flow. Number two -- I need the
call to be officially HIPAA-compliant. I know Teams is not covered under our BAA for
PHI discussions. We have been taking a risk. Number three -- I want to be able to
annotate images. Sometimes I am reviewing an X-ray with a patient and I want to draw
on it. I cannot do that in Teams.

[10:40] JORDAN KIM: As a patient, the joining process needs to be dead simple. I
tried to do a telehealth appointment with a different provider last year and I spent
12 minutes trying to figure out how to get into the call. By the time I got in I
was flustered. Can we just have a link in a text message?

[10:41] AMIR HOSSEINI: Absolutely. The patient flow will be: (1) you get an SMS
15 minutes before with a single-tap join link, (2) you tap it, it opens the browser
or app, (3) you are in the waiting room in under 30 seconds. No download required
for web. No account password needed if you have already verified via the SMS code.

[10:43] JORDAN KIM: And if the doctor is running late?

[10:44] AMIR HOSSEINI: The waiting room will show an estimated wait time updated in
real time, and if wait exceeds 15 minutes you automatically get an SMS so you can
stay on your phone rather than sitting at your desk staring at a loading screen.

[10:45] MARCUS WEBB: For billing purposes -- we need a timestamp of when the call
actually started and ended, not just when the room was opened. The clinical time
is what we bill. Can CareFlow provide that?

[10:47] AMIR HOSSEINI: Yes, we will log provider-joined-at, patient-joined-at,
session-ended-at, and duration in seconds. That is in the data model.

[10:48] DR. RAJIV PATEL: What about recordings? We sometimes want to record sessions
for quality review or for the patient's benefit.

[10:49] AMIR HOSSEINI: Recording requires explicit patient consent captured in the
system -- a digital checkbox with timestamp and IP. If consent is not given, the
recording button is disabled. Recordings go to S3 encrypted at rest. Access is
logged every time. We will not allow recording without consent -- it is a hard block.

---------------------------------------------------------------
SECTION 3: COMPLIANCE & SECURITY
---------------------------------------------------------------

[11:30] PRIYA SHARMA: Let us talk compliance. Carlos, you had flagged some specific
concerns from the Q1 audit findings.

[11:31] CARLOS DIAZ: Yes. Three findings from the Q1 audit. First: providers were
messaging patient PHI via Microsoft Teams which is not covered under our BAA. That
is the $45,000 fine. We need all clinical messaging to go through a system that is
under our HIPAA umbrella. Second: we had no centralized audit trail. Each system
kept its own log in different formats -- impossible to correlate. Third: session
management was weak -- no automatic timeout on the PHP portal, and some providers
share login credentials which is obviously terrible.

[11:35] AMIR HOSSEINI: CareFlow addresses all three. Secure in-app messaging under
the Twilio HIPAA BAA. Centralized Elasticsearch audit log -- every PHI access event,
every action, correlated by user and session. And hard 15-minute idle timeout with
MFA required for all clinical roles. Shared credentials are structurally impossible
-- each user has a unique account.

[11:38] CARLOS DIAZ: Good. One more thing -- data residency. All PHI must stay on
US soil. No data can transit outside the US even in a CDN cache.

[11:39] AMIR HOSSEINI: Confirmed. AWS deployment will be us-west-2 (primary) and
us-east-1 (DR). CloudFront will be configured with geo-restriction to US only.
No PHI will be cached at edge -- only static assets.

---------------------------------------------------------------
ACTION ITEMS
---------------------------------------------------------------

AI-01  Carlos Diaz      Obtain Athena API v2 sandbox credentials         By June 27, 2026
AI-02  Sarah Chen       Circulate draft BRD v2.0 for stakeholder review  By May 20, 2026
AI-03  Amir Hosseini    Confirm Twilio HIPAA BAA process timeline         By May 17, 2026
AI-04  Priya Sharma     Share telehealth UX wireframes (low-fi)           By May 21, 2026
AI-05  Lisa Tanaka      Provide sample scheduling data (3 months)         By May 22, 2026
AI-06  Marcus Webb      Share billing timestamp requirements doc           By May 19, 2026
AI-07  Priya Sharma     Schedule Workshop 4 -- Data Migration             Week of May 26, 2026

---------------------------------------------------------------
NEXT MEETING: Workshop 4 -- Data Migration & Integration
Date: May 27, 2026, 10:00 AM PST
---------------------------------------------------------------

[END OF TRANSCRIPT]
Transcribed by: Diego Ruiz, Business Analyst, Stackular Technologies
Reviewed by: Sarah Chen, MediConnect VP Engineering
"""
    with open(os.path.join(OUT, "MeetingTranscript_Workshop3_May14.txt"), "w") as f:
        f.write(content)
    print("  MeetingTranscript_Workshop3_May14.txt")

# =============================================================================
# 8. TRD — PDF
# =============================================================================

def gen_trd_pdf():
    doc, styles = pdf_doc("TRD_CareFlow_v1.0.pdf", "Technical Requirements Document")
    story = []
    story += [pdf_h(styles, "TECHNICAL REQUIREMENTS DOCUMENT", "Title"),
              pdf_h(styles, "CareFlow — Patient & Telehealth Management Platform", "h2"),
              pdf_p(styles, "Document ID: TRD-MC-2026-004  |  Version: 1.0  |  Date: June 3, 2026  |  Author: Carlos Diaz / Amir Hosseini"),
              pdf_sp(20)]

    story += [pdf_h(styles, "1. Approved Technology Stack")]
    stack = [["Layer", "Technology", "Rationale"],
             ["Frontend Web", "React 18 + TypeScript + Tailwind CSS", "Team familiarity; strong accessibility tooling"],
             ["Mobile", "React Native 0.74 (Expo)", "Code sharing with web; single JS team covers both platforms"],
             ["API Gateway", "Node.js 20 LTS + Express", "Thin proxy/auth layer; fast response times for scheduling"],
             ["Core Services", "Python 3.12 + FastAPI", "Scheduling engine, reminder engine -- async-first"],
             ["Database", "PostgreSQL 16 (AWS RDS Multi-AZ)", "ACID guarantees; row-level security for PHI"],
             ["Cache", "Redis 7.2 (AWS ElastiCache)", "Session store, slot reservation locks, rate limiting"],
             ["Search / Audit", "Elasticsearch 8.13", "Audit log (write-once), patient search, waitlist queries"],
             ["Telehealth", "Twilio Programmable Video", "HIPAA BAA available; WebRTC abstraction; recording API"],
             ["SMS", "Twilio SMS", "HIPAA BAA; reliable delivery; webhook support"],
             ["Email", "SendGrid", "Transactional email; template management"],
             ["Object Storage", "AWS S3 (SSE-AES256)", "Recording and document storage; lifecycle policies"],
             ["CDN", "AWS CloudFront (US-only)", "Static asset delivery; no PHI cached at edge"],
             ["Infrastructure", "AWS us-west-2 / us-east-1 DR", "MediConnect Enterprise Agreement; data residency"],
             ["IaC", "Terraform 1.8 + AWS CDK", "Reproducible environments; GitOps-compatible"],
             ["CI/CD", "GitHub Actions + AWS CodeDeploy", "Automated test > staging > production pipeline"],
             ["Monitoring", "CloudWatch + Datadog + PagerDuty", "Full-stack observability; on-call alerting"],
             ["SAST", "Snyk + Semgrep", "Automated vulnerability scanning on every PR"]]
    story += [pdf_table(stack, [1.5*inch, 2.2*inch, 2.8*inch]), pdf_sp()]

    story += [pdf_h(styles, "2. Security Architecture")]
    sec = [["Control", "Implementation"],
           ["Encryption at rest", "AES-256 on RDS, S3, ElastiCache (Redis AUTH + TLS)"],
           ["Encryption in transit", "TLS 1.3 minimum on all external and internal endpoints"],
           ["Field-level encryption", "Sensitive fields (SSN, insurance_id) encrypted via AWS KMS"],
           ["Authentication", "JWT RS256 + 15-min access tokens + rotating 8-hr refresh tokens"],
           ["MFA", "TOTP RFC 6238 enforced for all Provider, Admin, Compliance roles"],
           ["RBAC", "PostgreSQL row-level security policies enforce PHI access at DB layer"],
           ["Secrets management", "AWS Secrets Manager; no secrets in env vars or source code"],
           ["Pen testing", "Third-party OWASP-based pen test before launch (Week 26); annual thereafter"],
           ["HIPAA BAAs", "Executed with: AWS, Twilio, SendGrid before any PHI flows through those services"]]
    story += [pdf_table(sec, [2.5*inch, 4*inch]), pdf_sp()]

    story += [pdf_h(styles, "3. Performance Targets")]
    perf = [["Metric", "Target"],
            ["Patient portal page load (p95)", "< 2 seconds at 10,000 concurrent users"],
            ["API response time (p99)", "< 500ms for all non-streaming endpoints"],
            ["Telehealth call setup", "< 3 seconds from Join click to live video"],
            ["Reminder dispatch throughput", "5,000 SMS + 5,000 email per minute burst"],
            ["Appointment availability query", "< 200ms across 47 clinics"],
            ["Auto-scaling headroom", "Sustains 5x average peak without manual intervention"]]
    story += [pdf_table(perf, [3*inch, 3.5*inch]), pdf_sp()]

    story += [pdf_h(styles, "4. Data Retention & DR"),
              pdf_p(styles, "PHI retained 7 years (45 CFR s164.530). RDS: daily backups, 35-day retention, PITR to within 5 minutes. S3: lifecycle to Glacier after 90 days; delete after 7 years. DR: RDS read replica in us-east-1. RTO 4 hours, RPO 1 hour."),
              pdf_sp()]

    doc.build(story)
    print("  TRD_CareFlow_v1.0.pdf")

# =============================================================================
# 9. Data Dictionary — CSV
# =============================================================================

def gen_data_dict_csv():
    rows = [
        ["Table", "Column", "Type", "Nullable", "Description", "PII/PHI", "Retention"],
        ["patient", "id", "UUID", "N", "Primary key", "N", "7 years"],
        ["patient", "mrn", "VARCHAR(20)", "N", "Medical Record Number (auto-generated)", "Y", "7 years"],
        ["patient", "first_name", "VARCHAR(100)", "N", "Patient legal first name", "Y", "7 years"],
        ["patient", "last_name", "VARCHAR(100)", "N", "Patient legal last name", "Y", "7 years"],
        ["patient", "date_of_birth", "DATE", "N", "Patient date of birth", "Y", "7 years"],
        ["patient", "email", "VARCHAR(255)", "Y", "Contact email (encrypted at rest)", "Y", "7 years"],
        ["patient", "phone_mobile", "VARCHAR(20)", "Y", "Mobile phone (encrypted at rest)", "Y", "7 years"],
        ["patient", "preferred_language", "VARCHAR(10)", "N", "ISO 639-1 language code (default: en)", "N", "7 years"],
        ["patient", "portal_account_id", "UUID", "Y", "FK to portal_account.id (null if unregistered)", "N", "7 years"],
        ["patient", "created_at", "TIMESTAMPTZ", "N", "Record creation timestamp (UTC)", "N", "7 years"],
        ["patient", "updated_at", "TIMESTAMPTZ", "N", "Last update timestamp (UTC)", "N", "7 years"],
        ["provider", "id", "UUID", "N", "Primary key", "N", "7 years"],
        ["provider", "npi", "VARCHAR(10)", "N", "National Provider Identifier (unique)", "Y", "7 years"],
        ["provider", "first_name", "VARCHAR(100)", "N", "Provider legal first name", "Y", "7 years"],
        ["provider", "last_name", "VARCHAR(100)", "N", "Provider legal last name", "Y", "7 years"],
        ["provider", "specialty", "VARCHAR(100)", "N", "Clinical specialty", "N", "7 years"],
        ["provider", "clinic_id", "UUID", "N", "FK to clinic.id", "N", "7 years"],
        ["provider", "active", "BOOLEAN", "N", "Whether provider is currently active", "N", "7 years"],
        ["appointment", "id", "UUID", "N", "Primary key", "N", "7 years"],
        ["appointment", "patient_id", "UUID", "N", "FK to patient.id", "N", "7 years"],
        ["appointment", "provider_id", "UUID", "N", "FK to provider.id", "N", "7 years"],
        ["appointment", "clinic_id", "UUID", "N", "FK to clinic.id", "N", "7 years"],
        ["appointment", "service_type", "VARCHAR(50)", "N", "NEW_PATIENT | FOLLOW_UP | TELEHEALTH", "N", "7 years"],
        ["appointment", "status", "VARCHAR(20)", "N", "PENDING | CONFIRMED | ARRIVED | COMPLETED | NO_SHOW | CANCELLED", "N", "7 years"],
        ["appointment", "slot_start", "TIMESTAMPTZ", "N", "Appointment start (UTC)", "N", "7 years"],
        ["appointment", "slot_end", "TIMESTAMPTZ", "N", "Appointment end (UTC)", "N", "7 years"],
        ["appointment", "channel", "VARCHAR(20)", "N", "IN_PERSON | TELEHEALTH", "N", "7 years"],
        ["appointment", "reminder_48h_sent_at", "TIMESTAMPTZ", "Y", "Timestamp T-48h reminder dispatched", "N", "7 years"],
        ["appointment", "reminder_2h_sent_at", "TIMESTAMPTZ", "Y", "Timestamp T-2h reminder dispatched", "N", "7 years"],
        ["appointment", "cancelled_at", "TIMESTAMPTZ", "Y", "Cancellation timestamp", "N", "7 years"],
        ["appointment", "cancelled_by", "VARCHAR(20)", "Y", "PATIENT | PROVIDER | ADMIN", "N", "7 years"],
        ["telehealth_session", "id", "UUID", "N", "Primary key", "N", "7 years"],
        ["telehealth_session", "appointment_id", "UUID", "N", "FK to appointment.id", "N", "7 years"],
        ["telehealth_session", "twilio_room_sid", "VARCHAR(100)", "N", "Twilio Room SID", "N", "7 years"],
        ["telehealth_session", "provider_joined_at", "TIMESTAMPTZ", "Y", "Provider join timestamp", "N", "7 years"],
        ["telehealth_session", "patient_joined_at", "TIMESTAMPTZ", "Y", "Patient join timestamp", "N", "7 years"],
        ["telehealth_session", "ended_at", "TIMESTAMPTZ", "Y", "Session end timestamp", "N", "7 years"],
        ["telehealth_session", "duration_secs", "INTEGER", "Y", "Billable session duration in seconds", "N", "7 years"],
        ["telehealth_session", "recording_consent", "BOOLEAN", "N", "True if patient consented to recording", "N", "7 years"],
        ["telehealth_session", "recording_s3_key", "VARCHAR(500)", "Y", "S3 object key (null if not recorded)", "N", "7 years"],
        ["audit_event", "id", "UUID", "N", "Primary key", "N", "7 years"],
        ["audit_event", "user_id", "UUID", "N", "FK to portal_account.id", "N", "7 years"],
        ["audit_event", "user_role", "VARCHAR(30)", "N", "Role at time of action", "N", "7 years"],
        ["audit_event", "action", "VARCHAR(50)", "N", "READ | CREATE | UPDATE | DELETE | LOGIN | LOGOUT | EXPORT", "N", "7 years"],
        ["audit_event", "resource_type", "VARCHAR(50)", "N", "patient | appointment | message | recording", "N", "7 years"],
        ["audit_event", "resource_id", "UUID", "Y", "ID of affected resource", "N", "7 years"],
        ["audit_event", "ip_address", "VARCHAR(45)", "N", "IPv4/IPv6 of request origin", "N", "7 years"],
        ["audit_event", "session_id", "UUID", "N", "FK to active_session.id", "N", "7 years"],
        ["audit_event", "occurred_at", "TIMESTAMPTZ", "N", "Event timestamp (UTC, immutable)", "N", "7 years"],
        ["audit_event", "metadata", "JSONB", "Y", "Additional context (search query, export filters)", "N", "7 years"],
    ]
    with open(os.path.join(OUT, "DataDictionary_CareFlow_v1.0.csv"), "w", newline="") as f:
        w = csv.writer(f)
        w.writerows(rows)
    print("  DataDictionary_CareFlow_v1.0.csv")

# =============================================================================
# 10. Compliance Requirements — PDF
# =============================================================================

def gen_compliance_pdf():
    doc, styles = pdf_doc("ComplianceRequirements_HIPAA_SOC2.pdf", "Compliance Requirements")
    story = []
    story += [pdf_h(styles, "COMPLIANCE REQUIREMENTS DOCUMENT", "Title"),
              pdf_h(styles, "HIPAA & SOC 2 — CareFlow Platform", "h2"),
              pdf_p(styles, "Document ID: COMP-MC-2026-006  |  Version: 1.0  |  Author: James Okafor (Security Lead)  |  Date: June 4, 2026"),
              pdf_sp(20)]

    story += [pdf_h(styles, "1. Regulatory Framework"),
              pdf_p(styles, "CareFlow processes Protected Health Information (PHI) as defined under HIPAA (45 CFR Parts 160 and 164). MediConnect is a Covered Entity. Stackular Technologies acts as a Business Associate. BAAs must be in place with all sub-processors handling PHI before any PHI flows through those systems. CareFlow will also pursue SOC 2 Type II certification within 15 months of go-live."),
              pdf_sp()]

    story += [pdf_h(styles, "2. HIPAA Security Rule Requirements (45 CFR s164.312)")]
    hipaa = [["Control", "Specification", "CareFlow Implementation"],
             ["Access Control", "Unique user ID, auto-logoff, encryption", "RBAC + per-user accounts + 15-min timeout + AES-256"],
             ["Audit Controls", "Record and examine activity in PHI systems", "Elasticsearch write-once audit log for all PHI access"],
             ["Integrity", "Protect ePHI from improper alteration", "SHA-256 checksums; immutable audit log"],
             ["Person Authentication", "Verify identity before granting access", "Email/password + TOTP MFA for all clinical roles"],
             ["Transmission Security", "Protect ePHI over networks", "TLS 1.3 for all transmissions; no plaintext allowed"],
             ["Facility Access", "Physical safeguards for systems", "AWS data centres (SOC 2, ISO 27001 certified)"]]
    story += [pdf_table(hipaa, [1.8*inch, 2*inch, 2.7*inch]), pdf_sp()]

    story += [pdf_h(styles, "3. Required BAAs")]
    baa = [["Sub-processor", "PHI Exposure", "BAA Status", "Deadline"],
           ["AWS", "PHI in RDS, S3, ElastiCache", "Included in Enterprise Agreement", "Already executed"],
           ["Twilio (Video + SMS)", "Telehealth metadata + reminder content", "Twilio HIPAA BAA addendum required", "Must execute by Week 7"],
           ["SendGrid", "Email notifications (limited PHI)", "SendGrid HIPAA-compliant plan + BAA", "Must execute by Week 7"],
           ["Datadog", "Logs may contain PHI if unfiltered", "Datadog HIPAA BAA required", "Must execute by Week 10"]]
    story += [pdf_table(baa, [1.5*inch, 2*inch, 2*inch, 1*inch]), pdf_sp()]

    story += [pdf_h(styles, "4. Data Residency Requirements"),
              pdf_p(styles, "All PHI must remain on US soil at all times. Primary: AWS us-west-2. DR only: us-east-1. CloudFront configured with geo-restriction to US. No PHI may be cached at CDN edge nodes. Datadog log forwarding must route to US endpoints only."),
              pdf_sp()]

    story += [pdf_h(styles, "5. SOC 2 Trust Service Criteria (Target)")]
    soc = [["Criterion", "Requirement", "Target Date"],
           ["Security (CC)", "Logical and physical access controls; change management", "Audit kickoff Month 3 post-launch"],
           ["Availability (A)", "99.9% uptime SLA; monitoring; incident response", "Week 8"],
           ["Confidentiality (C)", "PHI encryption; access controls; BAAs", "Go-live"],
           ["Processing Integrity (PI)", "Complete and accurate processing; error handling", "Week 12"],
           ["Privacy (P)", "Notice; consent; collection; retention; disposal", "Week 16"]]
    story += [pdf_table(soc, [1.5*inch, 3*inch, 2*inch]), pdf_sp()]

    doc.build(story)
    print("  ComplianceRequirements_HIPAA_SOC2.pdf")

# =============================================================================
# 11. Structured Requirements JSON
# =============================================================================

def gen_requirements_json():
    data = {
        "project": "CareFlow",
        "client": CLIENT,
        "version": "1.0",
        "generated": "2026-06-05",
        "source_documents": ["BRD-MC-2026-002", "FRS-MC-2026-003", "SOW-MC-2026-001"],
        "epics": [
            {
                "id": "E-01",
                "title": "Patient Self-Service Portal",
                "description": "Enable patients to manage their healthcare journey digitally without front-desk assistance.",
                "user_stories": [
                    {
                        "id": "US-001", "title": "Patient Registration",
                        "as_a": "new patient",
                        "i_want": "to register for a portal account using my email or Google login",
                        "so_that": "I can access my health information online",
                        "acceptance_criteria": [
                            "AC-001-1: Registration form collects: first name, last name, date of birth, email, mobile",
                            "AC-001-2: Google SSO and Apple ID SSO both functional",
                            "AC-001-3: Verification email sent within 60 seconds",
                            "AC-001-4: Password breach-checked against HaveIBeenPwned before account creation"
                        ],
                        "story_points": 5, "priority": "Must"
                    },
                    {
                        "id": "US-002", "title": "Appointment Booking",
                        "as_a": "registered patient",
                        "i_want": "to book an appointment online in under 3 taps",
                        "so_that": "I do not need to call the clinic",
                        "acceptance_criteria": [
                            "AC-002-1: Patient can filter providers by specialty, clinic, and date",
                            "AC-002-2: Available slots shown in patient's local timezone",
                            "AC-002-3: Confirmation SMS and email sent within 2 minutes of booking",
                            "AC-002-4: Booking blocked if patient has 2+ uncancelled future appointments with same provider"
                        ],
                        "story_points": 8, "priority": "Must"
                    },
                    {
                        "id": "US-003", "title": "Appointment Reminders",
                        "as_a": "patient with an upcoming appointment",
                        "i_want": "to receive automated reminders",
                        "so_that": "I do not forget my appointment",
                        "acceptance_criteria": [
                            "AC-003-1: SMS reminder sent at exactly T-48h (+/- 5 minutes)",
                            "AC-003-2: Email reminder sent at exactly T-48h (+/- 5 minutes)",
                            "AC-003-3: SMS reminder sent at T-2h (+/- 5 minutes)",
                            "AC-003-4: Each reminder contains: provider name, date, time, location/join link, one-tap cancel link",
                            "AC-003-5: Patient can opt down to 1 reminder or opt out in profile settings"
                        ],
                        "story_points": 5, "priority": "Must"
                    },
                    {
                        "id": "US-004", "title": "Secure Messaging",
                        "as_a": "patient",
                        "i_want": "to message my care team securely",
                        "so_that": "I can communicate without calling the clinic",
                        "acceptance_criteria": [
                            "AC-004-1: Messages encrypted at rest and in transit",
                            "AC-004-2: Push notifications do not expose PHI in notification text",
                            "AC-004-3: Attachments: JPEG, PNG, PDF, DICOM only; max 25MB",
                            "AC-004-4: Message delivery confirmation shown to sender",
                            "AC-004-5: Messages retained 7 years; patient cannot delete messages"
                        ],
                        "story_points": 8, "priority": "Must"
                    }
                ]
            },
            {
                "id": "E-02",
                "title": "Provider Workflow Dashboard",
                "description": "Give providers a single interface for scheduling, patient records, and telehealth.",
                "user_stories": [
                    {
                        "id": "US-005", "title": "Daily Schedule View",
                        "as_a": "provider",
                        "i_want": "to see my full day schedule with patient details on one screen",
                        "so_that": "I can prepare for each appointment without switching systems",
                        "acceptance_criteria": [
                            "AC-005-1: Schedule loads within 2 seconds",
                            "AC-005-2: Each slot shows: patient name, MRN, service type, status, alerts",
                            "AC-005-3: Provider can filter by date and service type",
                            "AC-005-4: Status can be updated inline with single click"
                        ],
                        "story_points": 5, "priority": "Must"
                    },
                    {
                        "id": "US-006", "title": "Telehealth Session Launch",
                        "as_a": "provider",
                        "i_want": "to launch a telehealth session directly from the appointment entry",
                        "so_that": "I do not need to switch to a separate application",
                        "acceptance_criteria": [
                            "AC-006-1: 'Start Session' button visible on telehealth appointments from T-10min",
                            "AC-006-2: Clicking 'Start Session' creates Twilio room and opens video in same tab",
                            "AC-006-3: Patient intake form and prior notes accessible in split view during call",
                            "AC-006-4: Session duration timer visible throughout call",
                            "AC-006-5: Call quality indicator (green/amber/red) visible for both participants"
                        ],
                        "story_points": 13, "priority": "Must"
                    }
                ]
            },
            {
                "id": "E-03",
                "title": "Telehealth Platform",
                "description": "HIPAA-compliant WebRTC video consultations with no patient downloads required.",
                "user_stories": [
                    {
                        "id": "US-007", "title": "Patient Join Telehealth",
                        "as_a": "patient",
                        "i_want": "to join my telehealth appointment with a single tap from my SMS",
                        "so_that": "I can connect without technical difficulties",
                        "acceptance_criteria": [
                            "AC-007-1: SMS join link delivered 15 minutes before session",
                            "AC-007-2: Clicking link opens web browser; app opens if installed",
                            "AC-007-3: Patient in waiting room within 30 seconds of tapping link",
                            "AC-007-4: Waiting room shows provider name, estimated wait, connection test",
                            "AC-007-5: If wait > 15 minutes, patient receives SMS notification"
                        ],
                        "story_points": 8, "priority": "Must"
                    },
                    {
                        "id": "US-008", "title": "Session Recording with Consent",
                        "as_a": "provider",
                        "i_want": "to optionally record telehealth sessions with patient consent",
                        "so_that": "I can review consultations and share with patients",
                        "acceptance_criteria": [
                            "AC-008-1: Recording button disabled unless patient provides explicit digital consent",
                            "AC-008-2: Consent captured as checkbox with timestamp and IP address logged to audit",
                            "AC-008-3: Recording stored encrypted in S3; access logged",
                            "AC-008-4: Recording deleted after 7 years per retention policy"
                        ],
                        "story_points": 8, "priority": "Should"
                    }
                ]
            },
            {
                "id": "E-04",
                "title": "Compliance & Administration",
                "description": "Tools and controls to maintain HIPAA compliance and manage the platform.",
                "user_stories": [
                    {
                        "id": "US-009", "title": "Audit Log Access",
                        "as_a": "compliance officer",
                        "i_want": "to query all PHI access events",
                        "so_that": "I can demonstrate HIPAA compliance and investigate incidents",
                        "acceptance_criteria": [
                            "AC-009-1: Query filters: date range, user, patient MRN, action type, resource type",
                            "AC-009-2: Results returned in < 5 seconds for queries <= 90 days",
                            "AC-009-3: Exportable as CSV",
                            "AC-009-4: Audit records cannot be modified or deleted by any user",
                            "AC-009-5: Monthly compliance summary report auto-generated"
                        ],
                        "story_points": 8, "priority": "Must"
                    },
                    {
                        "id": "US-010", "title": "MFA Enforcement",
                        "as_a": "system",
                        "i_want": "to enforce MFA for all clinical and admin accounts",
                        "so_that": "unauthorized access to PHI is prevented",
                        "acceptance_criteria": [
                            "AC-010-1: TOTP MFA required on every login for Provider, Admin, Compliance, Billing roles",
                            "AC-010-2: Patients offered optional TOTP or SMS MFA",
                            "AC-010-3: No MFA bypass mechanism exists in production",
                            "AC-010-4: MFA setup mandatory at first login for clinical roles",
                            "AC-010-5: Lost MFA device recovery requires identity verification via IT admin"
                        ],
                        "story_points": 5, "priority": "Must"
                    }
                ]
            }
        ]
    }
    with open(os.path.join(OUT, "StructuredRequirements_CareFlow_v1.0.json"), "w") as f:
        json.dump(data, f, indent=2)
    print("  StructuredRequirements_CareFlow_v1.0.json")

# =============================================================================
# 12. Stakeholder Interview Notes — .txt
# =============================================================================

def gen_interview_notes():
    content = """STAKEHOLDER INTERVIEW NOTES -- CAREFLOW DISCOVERY
MediConnect Health Systems
Conducted by: Diego Ruiz & Priya Sharma (Stackular Technologies)
Period: April 28 - May 9, 2026
Total Interviews: 12 individuals across 7 roles

================================================================
INTERVIEW #1 -- Dr. Emily Rossetti, Provider (Portland Clinic)
Date: April 28, 2026 | Duration: 55 minutes
================================================================

Q: Walk me through your telehealth workflow today.
A: "It is painful. The patient books via the PHP portal. I do not get a calendar
invite -- I have to check the portal manually every morning. For the call, we use
personal Teams accounts. I know that is not compliant but there is no alternative.
During the call I have Athena open in another window but Teams keeps stealing focus.
I have missed entering notes mid-call multiple times."

Q: What is your biggest frustration with the current scheduling system?
A: "No-shows kill my day. I see maybe 24 patients a day. If 5 do not show, that is
5 time slots I could have filled with waitlisted patients. But our waitlist is a
printed paper at the front desk."

Q: What would the ideal telehealth experience look like?
A: "Patient arrives in a waiting room. I can see they are there. I click one button,
we are in the call. Their notes are on the right side of my screen. When we are done
I click end, write my note directly into the system, it syncs to Athena. Total time
to open and close an encounter: under 5 minutes."

KEY INSIGHTS:
  - Provider needs in-call access to patient chart (split view)
  - Real-time waitlist is high-value, not just nice-to-have
  - Change management/training as important as the software itself
  - Current Teams usage is acknowledged compliance risk

================================================================
INTERVIEW #2 -- Marcus Webb, Revenue Cycle Manager
Date: April 29, 2026 | Duration: 60 minutes
================================================================

Q: How does the no-show problem affect revenue cycle?
A: "We cannot bill for a no-show. A 30-min follow-up slot is $180 average
reimbursement. We lose 35 slots per day across the network. That is $6,300/day,
$1.57M/year. And that is before you account for the admin time chasing patients."

Q: What data do you need from the telehealth module for billing?
A: "Provider joined time, patient joined time, session end time, duration in minutes.
That is what we submit to insurance for a telehealth E&M code. Right now I am pulling
this from Teams meeting records, which is unreliable and a compliance mess."

Q: What reporting do you need from CareFlow?
A: "Daily: appointment counts by status broken down by clinic and provider. Weekly:
revenue at risk from upcoming appointments with patients who have >30% no-show
history. Monthly: trend data on no-show rate by service type. I need this as CSV
export or direct Power BI connection."

KEY INSIGHTS:
  - Precise session timestamps required (not just 'session occurred')
  - Power BI integration needed for revenue analytics
  - Marcus must be a UAT participant for billing-adjacent features
  - No-show history flag on patient record has direct revenue value

================================================================
INTERVIEW #3 -- Carlos Diaz, IT Infrastructure Lead
Date: April 30, 2026 | Duration: 75 minutes
================================================================

Q: What are your non-negotiables from an infrastructure standpoint?
A: "Data residency. Every byte of PHI stays in the US. We are AWS-native -- we have
an Enterprise Agreement. The solution must run on AWS. Third: I need runbooks for
every operational scenario -- failover, backup restoration, incident response."

Q: Walk me through your disaster recovery requirements.
A: "RTO 4 hours, RPO 1 hour. Board approved those targets last year. We had a 6-hour
outage in 2024 when the PHP portal's RDS ran out of storage. Providers could not book
appointments, front desk was writing on paper. I never want that again."

Q: Any concerns about the 32-week timeline?
A: "The security audit timeline worries me. Pen test in Week 26, go-live is Week 30.
If the pen test surfaces something serious, you have 4 weeks to fix it and retest.
That is tight. I would push for the pen test in Week 22 if possible."

KEY INSIGHTS:
  - US-only data residency is a hard disqualifier if violated
  - AWS mandatory (Enterprise Agreement)
  - Runbooks required as a formal deliverable
  - RTO 4h / RPO 1h board-mandated targets
  - Pen test timeline risk -- consider moving to Week 22

================================================================
INTERVIEW #4 -- Jordan Kim, Patient Representative
Date: May 6, 2026 | Duration: 45 minutes
================================================================

Q: What would make you actually use the patient portal?
A: "If I could see all my appointments in one place, message my doctor directly
instead of playing phone tag, and join a video call without downloading anything --
that alone would be a massive improvement."

Q: Telehealth experience -- what has gone wrong for you before?
A: "I spent 12 minutes trying to join a call once. I needed to download software,
then create an account, then it asked for a meeting ID which I did not have. By the
time I got in I was rattled. If the join experience is anything like that, patients
will not do it more than once."

KEY INSIGHTS:
  - No-download required for web telehealth is critical
  - Plain-language privacy explanation needed (not just legal boilerplate)
  - Simplicity of join flow is make-or-break for telehealth adoption

================================================================
CROSS-INTERVIEW THEMES (12 interviews total)
================================================================

THEME 1 -- Operational Fragmentation (11/12 interviews)
  All clinical stakeholders cited switching between systems as primary frustration.
  A unified interface is valued above any individual feature.

THEME 2 -- No-Show Economics (8/12 interviews)
  The $1.57M/year revenue impact is well understood. Automated reminder + waitlist
  is the most anticipated feature across all clinical interviews.

THEME 3 -- HIPAA Compliance Anxiety (7/12 interviews)
  The Q1 2026 fine created visible anxiety at all levels. Stakeholders want explicit
  assurance every channel is covered under a BAA, not just a verbal commitment.

THEME 4 -- Change Fatigue (6/12 interviews)
  Multiple stakeholders referenced the 2023 EHR migration as a painful precedent.
  Training quality and go-live support rated as highly as feature completeness.

THEME 5 -- Patient Experience Simplicity (4/12 interviews)
  Non-technical patients consistently prioritised simplicity over features.
  A single confusing step in the join flow will undermine telehealth adoption.

================================================================
[END OF INTERVIEW NOTES]
Prepared by: Diego Ruiz, Business Analyst, Stackular Technologies
Review date: May 12, 2026
================================================================
"""
    with open(os.path.join(OUT, "StakeholderInterviewNotes_May2026.txt"), "w") as f:
        f.write(content)
    print("  StakeholderInterviewNotes_May2026.txt")

# =============================================================================
# RUN ALL
# =============================================================================

if __name__ == "__main__":
    print(f"Generating documents into: {OUT}\n")
    gen_sow_docx()
    gen_sow_pdf()
    gen_brd_docx()
    gen_frs_docx()
    gen_rfp_pdf()
    gen_charter_docx()
    gen_transcript_txt()
    gen_trd_pdf()
    gen_data_dict_csv()
    gen_compliance_pdf()
    gen_requirements_json()
    gen_interview_notes()
    print("\nDone. All 12 documents generated.")
